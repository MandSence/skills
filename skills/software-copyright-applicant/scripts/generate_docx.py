# -*- coding: utf-8 -*-
"""
生成含目录的Word文档脚本

本模块用于将Markdown文档转换为Word(.docx)格式，
自动生成目录并插入图表图片。

优化内容：
- 支持识别并处理Mermaid转换失败的占位标记
- 在Word文档中显示转换失败的错误说明
- 确保图片插入位置与原始Markdown文档一致
"""

import re
import os
from typing import List, Dict, Any
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement

# 延迟导入，在命令行执行时动态加载
def _get_convert_functions():
    """动态获取Mermaid转换函数，避免模块导入问题"""
    import sys
    import importlib.util

    # 获取当前脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    convert_script = os.path.join(script_dir, 'convert_mermaid_to_png.py')

    # 动态加载模块
    spec = importlib.util.spec_from_file_location("convert_mermaid_to_png", convert_script)
    module = importlib.util.module_from_spec(spec)
    sys.modules["convert_mermaid_to_png"] = module
    spec.loader.exec_module(module)

    return module.convert_mermaid_to_png, module.replace_mermaid_blocks_with_images

# 初始化转换函数
_convert_mermaid_to_png = None
_replace_mermaid_blocks_with_images = None

def _ensure_convert_functions():
    """确保转换函数已加载"""
    global _convert_mermaid_to_png, _replace_mermaid_blocks_with_images
    if _convert_mermaid_to_png is None:
        _convert_mermaid_to_png, _replace_mermaid_blocks_with_images = _get_convert_functions()
    return _convert_mermaid_to_png, _replace_mermaid_blocks_with_images

def convert_mermaid_to_png(*args, **kwargs):
    """包装函数，用于调用convert_mermaid_to_png"""
    func, _ = _ensure_convert_functions()
    return func(*args, **kwargs)

def replace_mermaid_blocks_with_images(*args, **kwargs):
    """包装函数，用于调用replace_mermaid_blocks_with_images"""
    _, func = _ensure_convert_functions()
    return func(*args, **kwargs)


def set_chinese_font(run, font_name='宋体', size=12, bold=False, color=None):
    """
    设置中文字体

    Args:
        run: 文本运行对象
        font_name: 字体名称
        size: 字号（磅）
        bold: 是否加粗
        color: 颜色（RGBColor对象）
    """
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_styles(doc: Document):
    """
    设置文档样式

    Args:
        doc: Document对象
    """
    # 设置默认字体（正文）
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置标题1样式
    if 'Heading 1' in doc.styles:
        h1 = doc.styles['Heading 1']
        h1.font.name = '黑体'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        h1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 设置标题2样式
    if 'Heading 2' in doc.styles:
        h2 = doc.styles['Heading 2']
        h2.font.name = '黑体'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        h2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 设置标题3样式
    if 'Heading 3' in doc.styles:
        h3 = doc.styles['Heading 3']
        h3.font.name = '黑体'
        h3.font.size = Pt(13)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        h3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_page_numbering(doc: Document):
    """
    添加页码

    Args:
        doc: Document对象
    """
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加页码字段
    run = paragraph.add_run()
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def clean_markdown_text(text: str) -> str:
    """
    清理Markdown格式标记

    Args:
        text: 原始文本

    Returns:
        清理后的文本
    """
    # 移除四级标题标记 ####
    text = re.sub(r'^#{4,5,6}\s+', '', text, flags=re.MULTILINE)

    # 移除粗体标记 **text**
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)

    # 移除斜体标记 *text*
    text = re.sub(r'\*([^*]+)\*', r'\1', text)

    # 移除代码标记 `text`
    text = re.sub(r'`([^`]+)`', r'\1', text)

    # 移除链接标记 [text](url)，保留text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

    return text


def parse_markdown_with_toc(md_content: str) -> tuple:
    """
    解析Markdown文档，提取内容和生成目录条目

    Args:
        md_content: Markdown内容

    Returns:
        (内容列表, 目录条目列表)
        内容列表: [{'type': 'heading1'/'heading2'/'heading3'/'text'/'code'/'image'/'error_marker'/'table'/'list', 'content': ..., 'level': ...}, ...]
        目录条目列表: [{'level': 1/2/3, 'title': '...', 'page': 1}, ...]
    """
    lines = md_content.split('\n')
    content_items = []
    toc_entries = []
    page_counter = 1  # 假设目录占1页，从第2页开始

    i = 0
    toc_level_1_counter = 0

    # 匹配图表转换失败的占位标记：[图表X转换失败: ...]
    error_marker_pattern = re.compile(r'^\[图表(\d+)转换失败:\s*(.+)\]$')

    while i < len(lines):
        line = lines[i]

        # 解析一级标题 (#)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            content_items.append({
                'type': 'heading1',
                'content': title,
                'level': 1
            })
            toc_level_1_counter += 1
            toc_entries.append({
                'level': 1,
                'title': title,
                'numeral': get_chinese_numeral(toc_level_1_counter),
                'page': page_counter
            })
            page_counter += 1
            i += 1

        # 解析二级标题 (##)
        elif line.startswith('## ') and not line.startswith('###'):
            title = line[3:].strip()
            content_items.append({
                'type': 'heading2',
                'content': title,
                'level': 2
            })
            toc_entries.append({
                'level': 2,
                'title': title,
                'numeral': f'{toc_level_1_counter}.',
                'page': page_counter
            })
            page_counter += 1
            i += 1

        # 解析三级标题 (###)
        elif line.startswith('### ') and not line.startswith('####'):
            title = line[4:].strip()
            content_items.append({
                'type': 'heading3',
                'content': title,
                'level': 3
            })
            # 三级标题也加入目录
            toc_entries.append({
                'level': 3,
                'title': title,
                'numeral': f'{toc_level_1_counter}.',  # 与二级标题同级
                'page': page_counter
            })
            page_counter += 1
            i += 1

        # 跳过四级及以上标题（不加入目录，只作为正文）
        elif line.startswith('####'):
            title = clean_markdown_text(line.lstrip('#').strip())
            content_items.append({
                'type': 'heading4',
                'content': title,
                'level': 4
            })
            i += 1

        # 解析图表转换失败的占位标记
        elif line.strip().startswith('[') and ']转换失败:' in line:
            match = error_marker_pattern.match(line.strip())
            if match:
                chart_num = match.group(1)
                error_msg = match.group(2)
                content_items.append({
                    'type': 'error_marker',
                    'content': f'图表{chart_num}转换失败: {error_msg}',
                    'chart_num': chart_num
                })
                page_counter += 1
            else:
                # 不是错误标记，作为普通文本处理
                i += 1
                continue
            i += 1

        # 解析表格
        elif line.startswith('|'):
            # 收集整个表格
            table_lines = []
            while i < len(lines) and (lines[i].startswith('|') or lines[i].strip() == ''):
                if lines[i].strip():
                    table_lines.append(lines[i])
                i += 1

            if table_lines:
                content_items.append({
                    'type': 'table',
                    'content': table_lines
                })
                page_counter += 1
            continue

        # 解析代码块
        elif line.startswith('```'):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1

            if code_lines:
                content_items.append({
                    'type': 'code',
                    'content': '\n'.join(code_lines),
                    'lang': lang
                })
                page_counter += 1

        # 解析图片
        elif line.startswith('!['):
            # 提取图片路径
            match = re.match(r'!\[.*?\]\((.*?)\)', line)
            if match:
                img_path = match.group(1)
                content_items.append({
                    'type': 'image',
                    'content': img_path
                })
                page_counter += 1
            i += 1

        # 解析列表
        elif re.match(r'^\s*[-*+]\s+', line):
            # 收集整个列表
            list_lines = []
            while i < len(lines) and (re.match(r'^\s*[-*+]\s+', lines[i]) or lines[i].strip() == ''):
                if lines[i].strip():
                    list_lines.append(lines[i])
                i += 1

            if list_lines:
                content_items.append({
                    'type': 'list',
                    'content': list_lines
                })
                page_counter += 1
            continue

        # 普通文本段落
        elif line.strip():
            # 收集连续的文本行
            text_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith(('#', '|', '```', '!', '-', '*')):
                text_lines.append(lines[i])
                i += 1

            text = '\n'.join(text_lines)
            if text.strip():
                # 清理Markdown标记
                text = clean_markdown_text(text)
                content_items.append({
                    'type': 'text',
                    'content': text
                })
            continue

        else:
            i += 1

    return content_items, toc_entries


def get_chinese_numeral(n: int) -> str:
    """
    获取中文数字

    Args:
        n: 数字

    Returns:
        中文数字
    """
    numerals = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
                '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']
    if n <= len(numerals):
        return numerals[n - 1]
    return str(n)


def insert_toc(doc: Document, toc_entries: List[Dict]):
    """
    插入目录

    Args:
        doc: Document对象
        toc_entries: 目录条目列表
    """
    # 添加目录标题
    p = doc.add_paragraph()
    run = p.add_run('目  录')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    set_chinese_font(run, '黑体', 18, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)

    # 添加分隔线
    p = doc.add_paragraph()
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加目录条目（包括二级和三级标题）
    for entry in toc_entries:
        p = doc.add_paragraph()

        # 设置缩进
        if entry['level'] == 1:
            # 一级标题
            numeral = entry.get('numeral', get_chinese_numeral(toc_entries.index(entry) + 1))
            run = p.add_run(f'{numeral}、{entry["title"]}')
            p.paragraph_format.left_indent = Cm(0)
            set_chinese_font(run, '宋体', 12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            # 添加页码
            dots = '.' * (50 - len(entry['title']) - len(str(entry['page'])) - 10)
            p.add_run(f'{dots}{entry["page"]}')
            p.paragraph_format.line_spacing = 1.5
        elif entry['level'] == 2:
            # 二级标题
            numeral = entry.get('numeral', '')
            run = p.add_run(f'  {numeral}{entry["title"]}')
            p.paragraph_format.left_indent = Cm(1.5)
            set_chinese_font(run, '宋体', 12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            # 添加页码
            dots = '.' * (50 - len(entry['title']) - len(str(entry['page'])) - 10)
            p.add_run(f'{dots}{entry["page"]}')
            p.paragraph_format.line_spacing = 1.5
        elif entry['level'] == 3:
            # 三级标题
            run = p.add_run(f'    {entry["title"]}')
            p.paragraph_format.left_indent = Cm(2.5)
            set_chinese_font(run, '宋体', 12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            # 添加页码
            dots = '.' * (50 - len(entry['title']) - len(str(entry['page'])) - 10)
            p.add_run(f'{dots}{entry["page"]}')
            p.paragraph_format.line_spacing = 1.5

    # 添加分页符
    doc.add_page_break()


def insert_error_box(doc: Document, error_text: str):
    """
    在文档中插入错误提示框

    Args:
        doc: Document对象
        error_text: 错误文本
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)

    # 添加浅灰色背景
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'FFF3CD')  # 浅黄色背景，警告色
    p._element.get_or_add_pPr().append(shading_elm)

    # 添加边框
    p_bdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), '856404')
        p_bdr.append(border)
    p._element.get_or_add_pPr().append(p_bdr)

    # 添加图标和错误文本
    run = p.add_run('⚠ ')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(133, 100, 4)  # 深黄色文字

    run = p.add_run(error_text)
    set_chinese_font(run, '宋体', 12)
    run.font.color.rgb = RGBColor(133, 100, 4)  # 深黄色文字


def insert_content(doc: Document, content_items: List[Dict], assets_dir: str):
    """
    插入正文内容

    Args:
        doc: Document对象
        content_items: 内容列表
        assets_dir: 资源目录路径
    """
    for item in content_items:
        if item['type'] == 'heading1':
            p = doc.add_heading(item['content'], level=1)
            set_chinese_font(p.runs[0], '黑体', 16, bold=True)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # 黑色

        elif item['type'] == 'heading2':
            p = doc.add_heading(item['content'], level=2)
            set_chinese_font(p.runs[0], '黑体', 14, bold=True)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # 黑色

        elif item['type'] == 'heading3':
            p = doc.add_heading(item['content'], level=3)
            set_chinese_font(p.runs[0], '黑体', 13, bold=True)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # 黑色

        elif item['type'] == 'heading4':
            # 四级标题作为粗体正文处理
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(item['content'])
            set_chinese_font(run, '黑体', 12, bold=True)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色

        elif item['type'] == 'text':
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(item['content'])
            set_chinese_font(run, '宋体', 12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色

        elif item['type'] == 'code':
            p = doc.add_paragraph()
            # 添加浅灰色背景
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F5F5F5')
            p._element.get_or_add_pPr().append(shading_elm)

            run = p.add_run(item['content'])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)

        elif item['type'] == 'image':
            img_path = item['content']

            # 如果是相对路径，尝试在assets_dir中查找
            if not os.path.isabs(img_path):
                full_path = os.path.join(assets_dir, os.path.basename(img_path))
                if os.path.exists(full_path):
                    img_path = full_path

            if os.path.exists(img_path):
                try:
                    # 插入图片
                    doc.add_picture(img_path, width=Inches(6))
                    # 图片居中
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    last_paragraph.paragraph_format.space_before = Pt(12)
                    last_paragraph.paragraph_format.space_after = Pt(12)
                except Exception as e:
                    print(f"插入图片失败: {e}")
                    # 插入图片失败时，显示错误提示
                    insert_error_box(doc, f"图片插入失败: {str(e)}\n图片路径: {img_path}")

        elif item['type'] == 'error_marker':
            # 插入图表转换失败的错误提示框
            insert_error_box(doc, item['content'])

        elif item['type'] == 'table':
            # 解析表格
            table_lines = item['content']

            # 解析表头和数据行
            rows_data = []
            for line in table_lines:
                if line.strip().startswith('|'):
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    # 清理Markdown标记
                    cleaned_cells = [clean_markdown_text(cell) for cell in cells]
                    rows_data.append(cleaned_cells)

            if rows_data:
                # 创建表格
                table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                table.style = 'Light Grid Accent 1'

                for i, row_data in enumerate(rows_data):
                    row = table.rows[i]
                    for j, cell_data in enumerate(row_data):
                        cell = row.cells[j]
                        cell.text = cell_data
                        # 设置单元格文本格式
                        set_chinese_font(cell.paragraphs[0].runs[0], '宋体', 10.5)
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                        # 居中对齐
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 第一行加粗
                        if i == 0:
                            cell.paragraphs[0].runs[0].font.bold = True
                        # 设置行高
                        row.height = Cm(0.6)

                # 表格前后留白
                table.rows[0].height = Cm(0.8)

        elif item['type'] == 'list':
            # 解析列表
            list_lines = item['content']

            for line in list_lines:
                # 清理Markdown标记
                content = clean_markdown_text(re.sub(r'^\s*[-*+]\s+', '', line))
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(content)
                set_chinese_font(run, '宋体', 12)
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色


def generate_docx_with_toc(
    md_file: str,
    docx_file: str,
    assets_dir: str,
    software_name: str = None
):
    """
    生成含目录的Word文档

    Args:
        md_file: Markdown文件路径
        docx_file: 输出的Word文档路径
        assets_dir: 资源目录路径（用于查找图片）
        software_name: 软件名称（用于转换Mermaid图表）
    """
    # 转换为绝对路径并规范化
    md_file = os.path.abspath(md_file)
    docx_file = os.path.abspath(docx_file)
    assets_dir = os.path.abspath(assets_dir)

    # 检查文件是否存在
    if not os.path.exists(md_file):
        raise FileNotFoundError(f"Markdown文件不存在: {md_file}\n当前工作目录: {os.getcwd()}")

    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 转换Mermaid图表为PNG，并获取转换结果
    conversion_results = []
    if software_name:
        results = convert_mermaid_to_png(md_content, assets_dir, software_name)
        conversion_results = results
        # 替换Mermaid代码块为图片引用或占位标记
        md_content = replace_mermaid_blocks_with_images(md_content, results)

    # 解析Markdown内容
    content_items, toc_entries = parse_markdown_with_toc(md_content)

    # 创建Word文档
    doc = Document()

    # 设置样式
    set_styles(doc)

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 插入目录
    insert_toc(doc, toc_entries)

    # 插入正文内容
    insert_content(doc, content_items, assets_dir)

    # 添加页码
    add_page_numbering(doc)

    # 保存文档
    doc.save(docx_file)
    print(f"Word文档已生成: {docx_file}")

    # 输出转换结果摘要
    if conversion_results:
        failed_count = sum(1 for r in conversion_results if not r['success'])
        if failed_count > 0:
            print(f"注意: {failed_count} 个Mermaid图表转换失败，已在文档中标注错误信息。")


def generate_simple_docx(
    md_file: str,
    docx_file: str,
    assets_dir: str = None
):
    """
    生成简单的Word文档（不含目录，直接转换）

    Args:
        md_file: Markdown文件路径
        docx_file: 输出的Word文档路径
        assets_dir: 资源目录路径（可选，用于查找图片）
    """
    if assets_dir is None:
        assets_dir = os.path.dirname(md_file)

    # 使用空软件名，不会转换Mermaid
    generate_docx_with_toc(md_file, docx_file, assets_dir, software_name=None)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description='将Markdown文档转换为Word(.docx)格式')
    parser.add_argument('md_file', help='Markdown文件路径')
    parser.add_argument('docx_file', help='输出的Word文档路径')
    parser.add_argument('assets_dir', help='资源目录路径（用于查找和保存图片）')
    parser.add_argument('--software-name', help='软件名称（用于转换Mermaid图表，如果省略则不转换）')

    args = parser.parse_args()

    # 确保assets目录存在
    os.makedirs(args.assets_dir, exist_ok=True)

    # 生成Word文档
    generate_docx_with_toc(args.md_file, args.docx_file, args.assets_dir, args.software_name)
